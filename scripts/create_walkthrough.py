from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/VishGym_Solution_Walkthrough.docx")
BENCHMARK = Path("artifacts/benchmarks/adapter-smoke-v7/benchmark.json")
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5B6573"
LIGHT = "E8EEF5"
GRAY = "F2F4F7"


def set_run_font(run, size: float, color: str = "000000", bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_paragraph(paragraph, before: float = 0, after: float = 6, line: float = 1.10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_body(doc, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph(p)
    set_run_font(p.add_run(text), 11)


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(p, after=4, line=1.167)
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        set_run_font(p.add_run(item), 11)


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph(p, before={1: 16, 2: 12, 3: 8}[level], after={1: 8, 2: 6, 3: 4}[level])
    run = p.add_run(text)
    set_run_font(run, {1: 16, 2: 13, 3: 12}[level], {1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, line=1.0)
        set_run_font(p.add_run(text), 10, NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.10)
            set_run_font(p.add_run(text), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    for level, size, color in ((1, 16, BLUE), (2, 13, BLUE), (3, 12, DARK_BLUE)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(header, after=0)
    set_run_font(header.add_run("VISHGYM | CLOSED SYNTHETIC RESEARCH PROTOTYPE"), 8.5, MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(footer, after=0)
    set_run_font(footer.add_run("Synthetic-only demonstration. No real payments, identities, contacts, sites, or voices."), 8, MUTED)

    p = doc.add_paragraph()
    set_paragraph(p, before=4, after=4)
    set_run_font(p.add_run("MASTERCARD INNOVATION CHALLENGE 2026"), 10, BLUE, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, after=6)
    set_run_font(p.add_run("VishGym"), 26, NAVY, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, after=16)
    set_run_font(p.add_run("Multimodal Red/Blue Self-Play Arena for Payment-Fraud Defense"), 14, MUTED)

    add_table(doc, ["Prepared for", "Prototype focus", "Status"], [["Mastercard Innovation Challenge 2026", "Closed-loop AI defense lab for synthetic UPI social-engineering simulations", "Runnable prototype; measured Modal benchmark included"]], [2160, 5040, 2160])

    add_heading(doc, "1. Executive summary")
    add_body(doc, "VishGym is a closed, synthetic training arena that lets a Red agent and a Blue agent learn from each other without exposing real people or infrastructure to risk. The Red agent initiates fictional payment-fraud scenarios; the Blue agent uses only virtual tools to verify, refuse, report, or escalate. A fixed hybrid judge scores the complete episode and supplies delayed reward for alternating policy improvement.")
    add_body(doc, "The prototype is designed to satisfy the identify, generate, and defend loop: it maps a broad attack catalogue, creates repeatable multimodal simulations, and measures defensive policy performance against historical attacker checkpoints.")

    add_heading(doc, "2. Problem and innovation")
    add_table(doc, ["Challenge pillar", "VishGym response"], [
        ["Identify", "Nine high-level synthetic payment-fraud archetypes spanning voice, SMS, WhatsApp, support, invoice, refund, and cross-channel surfaces."],
        ["Generate", "Two role-specific agents produce audio-first, tool-using scenarios inside a fully isolated container."],
        ["Defend", "A Blue policy learns to choose safe actions while a fixed judge measures fraud decisions, false blocks, and simulated compromise outcomes."],
    ], [2340, 7020])

    add_heading(doc, "3. Closed synthetic architecture")
    add_body(doc, "The environment implements OpenEnv-style reset, step, and state semantics. Each episode creates new fictional personas, pseudo-identifiers, a virtual INR wallet, synthetic inbox records, sandbox-only pages, fixed search results, audio turns, and an immutable tool ledger. The web prototype defaults to a fresh random seed, persona pair, pseudo-identity packet, built-in voice, and tone profile for every live call; fixed seeds remain available for reproducible evaluation.")
    add_bullets(doc, [
        "Red and Blue use separate QLoRA adapters over a shared Gemma 4 E2B multimodal base model.",
        "Conversation is audio-only for the opponent. The prototype displays the actual synthetic message spoken by either agent for judge review, but the agents themselves never receive a text transcript.",
        "Qwen3-TTS CustomVoice renders English turns using built-in, non-identifiable speakers and style instructions. The public runtime accepts no reference audio.",
        "The fixed judge combines deterministic ledger rules with a bounded contextual review, preventing reward drift or unreviewed judge training.",
    ])

    add_heading(doc, "4. Sandbox tools and boundaries")
    add_table(doc, ["Role", "Synthetic tools", "Hard boundary"], [
        ["Blue", "Inbox read/report/block; virtual wallet balance/pay/decline; pseudo-credential view; sandbox browser; fixed-corpus search.", "No external accounts, recipients, URLs, uploads, or payments."],
        ["Red", "Synthetic message send; local portal-template creation; fixed-corpus search.", "No outbound messaging, exposed portals, or real web search."],
        ["Judge", "Read-only access to audio, hidden transcript, and action ledger after episode completion.", "Cannot take actions or change its own scoring policy."],
    ], [1260, 4680, 3420])

    add_heading(doc, "5. Live demonstration scenarios")
    add_bullets(doc, [
        "Synthetic vishing followed by a virtual UPI collect request.",
        "Synthetic SMS payment-link impersonation in a sandbox-only browser flow.",
        "Synthetic WhatsApp beneficiary or invoice-change impersonation.",
    ])
    add_body(doc, " Six additional catalogue cards support breadth without claiming unimplemented live flows. All scenario language, identifiers, and pages are fictional and labelled as simulation content.")

    add_heading(doc, "6. Rewards and learning loop")
    add_body(doc, "Rewards are delayed until the terminal episode. Deterministic rules govern virtual payment, pseudo-credential exposure, reporting, blocking, and invalid actions. The judge provides only a small contextual adjustment. Blue is rewarded for correct safe behavior and penalized for unsafe virtual actions, unnecessary blocks, and unnecessary declines on legitimate controls; this prevents a degenerate block-everything policy. Red receives simulated reward only for sandbox-contained outcomes and loses reward for invalid or boundary-violating actions.")
    add_bullets(doc, [
        "Warm start: synthetic tool-use traces validate schemas and sandbox behavior.",
        "Red update: QLoRA/GRPO against frozen Blue, gated by held-out scenarios and historical Blue checkpoints.",
        "Blue update: QLoRA/GRPO against reviewed Red, gated by all nine fraud cards, at least two held-out seeds, F1 >= 0.80, legitimate false-block rate <= 10%, valid tool-call rate >= 98%, and zero boundary violations.",
        "Release: human review of metrics, seeds, data revision, and model manifest; no automatic promotion.",
    ])

    doc.add_page_break()
    add_heading(doc, "7. Training, reproducibility, and deployment")
    add_body(doc, "Google Colab MCP is the intended training operator surface. The accompanying notebook pulls the repository, validates a deterministic local rollout, and records the required alternating-QloRA/GRPO workflow. Candidate adapters, synthetic datasets, metrics, seeds, and reviewer decisions are versioned as artifacts before promotion.")
    add_body(doc, "The prototype deploys as a Hugging Face GPU Docker Space: Streamlit provides the judge-facing UI, FastAPI exposes the simulator and model manifest, and Nginx routes both through port 7860. Runtime audio and simulation data are held in memory or ephemeral storage and expire after 60 minutes.")

    add_heading(doc, "8. Measured Modal benchmark")
    benchmark = json.loads(BENCHMARK.read_text(encoding="utf-8"))
    blue = benchmark["blue"]["report"]
    red = benchmark["red"]
    protocol = benchmark["protocol"]
    add_body(doc, "A completed Modal L4 benchmark evaluated two 2-step QLoRA validation adapters against each other using generated Qwen3-TTS CustomVoice audio. It covers all nine synthetic fraud cards plus one legitimate control at held-out seed 101. The system used audio-only opponent input, deterministic generation (temperature 0), and a 96-token action cap. The report stores aggregate metrics only; no raw audio, transcript, or model completion was persisted.")
    add_table(doc, ["Metric", "Observed result", "Interpretation"], [
        ["Blue fraud-decision F1", f"{blue['fraud_decision_f1']:.4f} (9/9 fraud cases safely defended)", "Strong fraud-card coverage in this one-seed validation run."],
        ["Legitimate false-block rate", f"{blue['legitimate_false_block_rate']:.0%} (1/1 control blocked)", "Fails the <=10% gate; adapter is not eligible for promotion."],
        ["Blue valid tool-call rate", f"{blue['valid_tool_call_rate']:.0%} ({blue['valid_tool_calls']}/{blue['total_tool_calls']})", "No malformed or boundary-violating tool action was observed."],
        ["Red simulated attack success", f"{red['simulated_attack_success_rate']:.0%} ({red['simulated_compromises']}/{red['total_episodes']})", "The validation attacker did not obtain a simulated compromise."],
        ["Red valid tool-call rate", f"{red['valid_tool_call_rate']:.0%} ({red['total_tool_calls']} calls)", "All attacker tool use remained inside the sandbox."],
    ], [2160, 2700, 4680])
    add_body(doc, "This is a reproducible pipeline benchmark, not a deployment-efficacy claim: it uses one held-out seed and undertrained adapters. The red/blue review manifests remain review_required. The next gate is a larger, multi-seed run (101 and 103), all nine fraud cards, held-out personas/timbres, legitimate controls, and human review; only then can the planned F1, false-block, validity, and boundary thresholds be assessed.")

    add_heading(doc, "9. Evaluation plan")
    add_bullets(doc, [
        "Hold out persona pairs, built-in timbres, scenario combinations, temperature values, noise, and latency settings.",
        "Report Blue fraud-decision F1, legitimate false-block rate, simulated compromise rate, valid tool-call rate, and gains over the initial Blue adapter.",
        "Compare promoted Blue versions against frozen historical Red checkpoints to detect overfitting to a single attacker policy.",
        "Reject narrow runs before review unless they cover all nine fraud attack cards across at least two held-out seeds.",
        "Run integration checks for invalid tool calls, non-sandbox URL attempts, inference timeouts, audio failures, and expired simulation records.",
    ])

    add_heading(doc, "10. Real-world feasibility and safeguards")
    add_body(doc, "VishGym is a training and evaluation environment, not a customer-facing fraud engine. In a production payment system, the Blue policy would become one risk signal alongside transaction controls, verified device and account signals, customer confirmation, and human escalation. It does not authorize or block real payments on its own.")
    add_body(doc, "The project deliberately forbids real-person voice retrieval, voice uploads, real identifiers, external browsing, live mailboxes, real payment rails, exposed portals, and automatic adapter publication. These constraints make the prototype suitable for defensive research while preserving a realistic closed-loop learning objective.")

    add_heading(doc, "11. Submission artifacts")
    add_bullets(doc, [
        "Runnable GitHub repository with environment, simulated tools, API, Streamlit UI, Docker deployment configuration, and unit tests.",
        "Colab notebook for reviewed training runs through the Google Colab MCP bridge.",
        "Hugging Face Docker Space configuration ready for deployment; no public hosted Space is claimed until it is deployed and reviewed.",
    ])

    doc.core_properties.title = "VishGym Solution Walkthrough"
    doc.core_properties.subject = "Closed synthetic multimodal red/blue self-play arena"
    doc.core_properties.author = "VishGym"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
